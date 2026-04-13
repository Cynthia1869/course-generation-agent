from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.core.schemas import SourceChunk, SourceDocument


class DocumentParser:
    def parse_file(self, path: Path, mime_type: str) -> SourceDocument:
        doc = SourceDocument(filename=path.name, mime_type=mime_type)
        try:
            text = self._extract_text(path, mime_type)
            chunks = [
                SourceChunk(text=chunk.strip(), index=index)
                for index, chunk in enumerate(self._chunk_text(text))
                if chunk.strip()
            ]
            doc.text_chunks = chunks
            doc.extract_status = "parsed"
            doc.metadata = {"path": str(path), "chunks": len(chunks)}
        except Exception as exc:  # noqa: BLE001
            doc.extract_status = "failed"
            doc.metadata = {"path": str(path), "error": str(exc)}
        return doc

    def _extract_text(self, path: Path, mime_type: str) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"} or mime_type.startswith("text/"):
            return path.read_text(encoding="utf-8")
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".docx":
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        raise ValueError(f"Unsupported file type: {suffix}")

    def _chunk_text(self, text: str, size: int = 1000) -> list[str]:
        return [text[i : i + size] for i in range(0, len(text), size)] or [text]
